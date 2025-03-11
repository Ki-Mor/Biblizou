"""
Auteur : ExEco Environnement - François Botcazou
Date de création : 2025/02
Dernière mise à jour : 2025/03
Version : 1.0
Nom : ZnieffXmlToDocx.py
Groupe : Biblizou_PatNat
Description : Module pour extraire les descriptions des sites ZNIEFF des fichiers XML ZNIEFF et créer un fichier de synthèse docx.
    Il est conçu pour être utilisé dans une extension QGIS
Dépendances :
    - Python 3.x
    - QGIS (QgsMessageLog)
    - xml.etree.ElementTree
    - python-docx
    - os, datetime

Utilisation :
    Ce module doit être appelé depuis une extension QGIS. Il prend en entrée un
    dossier contenant des fichiers XML et génère un fichier DOCX récapitulatif.
"""

from qgis.PyQt.QtWidgets import QFileDialog
from qgis.core import QgsMessageBar
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import os


class ZnieffXmlToDocx:
    def __init__(self, iface):
        self.iface = iface

    def run(self):
        """Exécuter le module."""
        try:
            folder_path = self.obtain_folder_path()
            if not folder_path:
                return
            self.process_xml_files_in_folder(folder_path)
            self.iface.messageBar().pushMessage("Succès", "Le fichier DOCX a été généré avec succès!",
                                                level=QgsMessageBar.INFO)
        except Exception as e:
            self.iface.messageBar().pushMessage("Erreur", str(e), level=QgsMessageBar.CRITICAL)

    def obtain_folder_path(self):
        """Ouvre un dialogue pour sélectionner un dossier contenant les fichiers XML."""
        folder = QFileDialog.getExistingDirectory(None, "Sélectionner un dossier contenant les fichiers XML")
        if not folder:
            self.iface.messageBar().pushMessage("Info", "Aucun dossier sélectionné.", level=QgsMessageBar.WARNING)
        return folder

    def xml_to_docx(self, xml_file, doc):
        """Extrait les données des fichiers XML et les ajoute au document DOCX."""
        try:
            tree = ET.parse(xml_file)
            root = tree.getroot()

            for znieff_elem in root.iter('ZNIEFF'):
                nm_sffzn_elem = znieff_elem.find('NM_SFFZN')
                lb_zn_elem = znieff_elem.find('LB_ZN')
                tx_gene_elem = znieff_elem.find('TX_GENE')

                lb_zn = lb_zn_elem.text if lb_zn_elem is not None else ''
                nm_sffzn = nm_sffzn_elem.text if nm_sffzn_elem is not None else ''
                combined_text = f"{lb_zn} - {nm_sffzn}"

                para = doc.add_paragraph()
                run = para.add_run(combined_text)
                run.bold = True
                run.underline = True
                run.font.color.rgb = RGBColor(0, 153, 153)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

                if tx_gene_elem is not None:
                    for p_elem in tx_gene_elem.findall('p'):
                        p_text = p_elem.text if p_elem is not None else ''
                        para = doc.add_paragraph(p_text)
                        para.paragraph_format.left_indent = Pt(28)
                        run = para.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                doc.add_paragraph('')

        except ET.ParseError as e:
            self.iface.messageBar().pushMessage("Erreur", f"Erreur d'analyse XML dans {xml_file}: {e}",
                                                level=QgsMessageBar.CRITICAL)
        except Exception as e:
            self.iface.messageBar().pushMessage("Erreur", f"Erreur avec le fichier {xml_file}: {e}",
                                                level=QgsMessageBar.CRITICAL)

    def clean_document(self, doc):
        """Nettoie le document en supprimant les espaces et lignes vides inutiles."""
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = ' '.join(run.text.split())
            if para.text.strip() == '':
                para.clear()

    def process_xml_files_in_folder(self, folder_path):
        """Traite tous les fichiers XML d'un dossier et génère un fichier DOCX."""
        if not os.path.isdir(folder_path):
            self.iface.messageBar().pushMessage("Erreur", f"Le chemin {folder_path} n'est pas un dossier valide.",
                                                level=QgsMessageBar.CRITICAL)
            return

        xml_files = [f for f in os.listdir(folder_path) if
                     f.endswith('.xml') and not f.startswith('FR') and len(f) == 13]
        if not xml_files:
            self.iface.messageBar().pushMessage("Info", "Aucun fichier XML valide trouvé dans le dossier.",
                                                level=QgsMessageBar.WARNING)
            return

        current_time = datetime.now().strftime("%Y%m%d%H%M%S")
        docx_file = os.path.join(folder_path, f'ZNIEFF_Descriptions_des_sites_{current_time}.docx')
        doc = Document()

        for xml_file in xml_files:
            full_path = os.path.join(folder_path, xml_file)
            self.xml_to_docx(full_path, doc)

        self.clean_document(doc)
        doc.save(docx_file)
        self.iface.messageBar().pushMessage("Succès", f"Document créé : {docx_file}", level=QgsMessageBar.INFO)
