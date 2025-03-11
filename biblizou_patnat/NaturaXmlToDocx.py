"""
Auteur : ExEco Environnement - François Botcazou
Date de création : 2025/02
Dernière mise à jour : 2025/03
Version : 1.0
Nom : NaturaXmlToDocx.py
Groupe : Biblizou_PatNat
Description : Module pour extraire les descriptions des sites ZNIEFF des fichiers XML ZNIEFF et créer un fichier de synthèse docx.
    Il est conçu pour être utilisé dans une extension QGIS
Dépendances :
    - Python 3.x
    - QGIS (QgsMessageBar, QgsMessageLog)
    - xml.etree.ElementTree
    - python-docx
    - os, datetime, PyQt5.QtWidgets

Utilisation :
    Ce module doit être appelé depuis une extension QGIS. Il prend en entrée un
    dossier contenant des fichiers XML et génère un fichier DOCX récapitulatif.
"""

import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import os
from PyQt5.QtWidgets import QFileDialog, QMessageBox
from qgis.core import QgsMessageBar, QgsProject

class NaturaXmlToDocx:
    def __init__(self, iface):
        self.iface = iface

    def obtain_folder_path(self):
        """ Ouvre une boîte de dialogue pour sélectionner un dossier """
        folder_path = QFileDialog.getExistingDirectory(None, "Sélectionner un dossier contenant les fichiers XML")
        if not folder_path:
            self.iface.messageBar().pushMessage("Annulation", "Aucun dossier sélectionné", level=QgsMessageBar.WARNING)
        return folder_path

    def xml_to_docx(self, xml_file, doc):
        try:
            tree = ET.parse(xml_file)
            root = tree.getroot()
            for n2000_elem in root.iter('BIOTOP'):
                sitecode_elem = n2000_elem.find('SITECODE')
                site_name_elem = n2000_elem.find('SITE_NAME')
                site_name = site_name_elem.text if site_name_elem is not None else ''
                sitecode = sitecode_elem.text if sitecode_elem is not None else ''
                combined_text = f"{site_name} - {sitecode}"
                para = doc.add_paragraph()
                run = para.add_run(combined_text)
                run.bold = True
                run.underline = True
                run.font.color.rgb = RGBColor(0, 153, 153)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                commentaire_elem = n2000_elem.find('COMMENTAIRE')
                if commentaire_elem is not None:
                    for commentaire_row_elem in commentaire_elem.iter('COMMENTAIRE_ROW'):
                        for tag in ['QUALITY', 'VULNAR']:
                            elem = commentaire_row_elem.find(tag)
                            if elem is not None and elem.text:
                                para = doc.add_paragraph(elem.text.strip())
                                para.paragraph_format.left_indent = Pt(28)
                                run = para.runs[0]
                                run.font.name = 'Calibri'
                                run.font.size = Pt(11)
                                run.font.color.rgb = RGBColor(0, 0, 0)
                doc.add_paragraph('')
        except ET.ParseError as e:
            self.iface.messageBar().pushMessage("Erreur XML", f"Erreur de parsing dans {xml_file}: {e}", level=QgsMessageBar.CRITICAL)
        except Exception as e:
            self.iface.messageBar().pushMessage("Erreur", f"Erreur inattendue avec {xml_file}: {e}", level=QgsMessageBar.CRITICAL)

    def clean_document(self, doc):
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = ' '.join(run.text.split())
            if para.text.strip() == '':
                para.clear()

    def process_xml_files_in_folder(self, folder_path):
        if not os.path.isdir(folder_path):
            self.iface.messageBar().pushMessage("Erreur", f"Le chemin {folder_path} n'est pas un répertoire valide.", level=QgsMessageBar.CRITICAL)
            return
        xml_files = [f for f in os.listdir(folder_path) if f.startswith('FR') and f.endswith('.xml') and len(f) == 13]
        if not xml_files:
            self.iface.messageBar().pushMessage("Information", "Aucun fichier XML trouvé dans le dossier.", level=QgsMessageBar.WARNING)
            return
        current_time = datetime.now().strftime("%Y%m%d%H%M%S")
        docx_file = os.path.join(folder_path, f'N2000_Descriptions_des_sites_{current_time}.docx')
        doc = Document()
        for xml_file in xml_files:
            full_path = os.path.join(folder_path, xml_file)
            self.xml_to_docx(full_path, doc)
        self.clean_document(doc)
        try:
            doc.save(docx_file)
            self.iface.messageBar().pushMessage("Succès", f"Document créé : {docx_file}", level=QgsMessageBar.INFO)
        except Exception as e:
            self.iface.messageBar().pushMessage("Erreur", f"Erreur lors de l'enregistrement du document: {e}", level=QgsMessageBar.CRITICAL)

    def run(self):
        folder_path = self.obtain_folder_path()
        if folder_path:
            self.process_xml_files_in_folder(folder_path)