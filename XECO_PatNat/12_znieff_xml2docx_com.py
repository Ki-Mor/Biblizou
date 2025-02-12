"""
Author : ExEco Environnement
Edition date : 2025/02
Name : 12_znieff_xml2docx_com
Group : Biblio_PatNat
"""

import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import os
import sys

def obtain_folder_path():
    return sys.argv[1] if len(sys.argv) > 1 else input("Entrez le chemin du dossier contenant les fichiers XML : ")

def xml_to_docx(xml_file, doc):
    """
    Parse the XML file and append the concatenated <LB_ZN>, <NM_SFFZN>,
    followed by a line break and the content inside each <p> tag in <TX_GENE> from each <ZNIEFF> element.
    """
    try:
        # Parse the XML file
        tree = ET.parse(xml_file)
        root = tree.getroot()

        # Iterate through <ZNIEFF> elements
        for znieff_elem in root.iter('ZNIEFF'):
            # Extract NM_SFFZN and LB_ZN from the current <ZNIEFF> element
            nm_sffzn_elem = znieff_elem.find('NM_SFFZN')
            lb_zn_elem = znieff_elem.find('LB_ZN')
            tx_gene_elem = znieff_elem.find('TX_GENE')

            # Initialize variables for LB_ZN and NM_SFFZN
            lb_zn = lb_zn_elem.text if lb_zn_elem is not None else ''
            nm_sffzn = nm_sffzn_elem.text if nm_sffzn_elem is not None else ''
            tx_gene = tx_gene_elem.text if tx_gene_elem is not None else ''

            # Concatenate LB_ZN and NM_SFFZN
            combined_text = f"{lb_zn} - {nm_sffzn}"

            # Add the concatenated text as a paragraph in the DOCX document (with formatting)
            para = doc.add_paragraph()
            run = para.add_run(combined_text)
            run.bold = True
            run.underline = True
            run.font.color.rgb = RGBColor(0, 153, 153)  # Color #009999
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

            # Check if TX_GENE exists and contains <p> elements
            if tx_gene_elem is not None:
                # Iterate over all <p> tags inside <TX_GENE> and extract their text
                for p_elem in tx_gene_elem.findall('p'):
                    p_text = p_elem.text if p_elem is not None else ''
                    # Add the content of <p> as a paragraph in the DOCX document (with formatting)
                    para = doc.add_paragraph(p_text)
                    para.paragraph_format.left_indent = Pt(28)  # 1 cm tabulation
                    run = para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

            # Add a line break between entries (if needed)
            doc.add_paragraph('')

    except ET.ParseError as e:
        print(f"Error parsing XML file {xml_file}: {e}")
    except Exception as e:
        print(f"Unexpected error with file {xml_file}: {e}")

def clean_document(doc):
    """
    Remove double spaces and double line breaks from the document, preserving the formatting.
    """
    for para in doc.paragraphs:
        # Iterate over the runs in the paragraph
        for run in para.runs:
            # Clean the text of the current run (remove double spaces)
            run.text = ' '.join(run.text.split())

        # Remove unnecessary empty paragraphs (double line breaks)
        if para.text.strip() == '':
            para.clear()  # Clears the paragraph content, removing extra empty lines

def process_xml_files_in_folder(folder_path):
    """
    Process all XML files in a folder and append their extracted data into a single DOCX file.
    """
    # Validate the folder path
    if not os.path.isdir(folder_path):
        print(f"The path {folder_path} is not a valid directory.")
        return

    # List all XML files in the folder
    xml_files = [f for f in os.listdir(folder_path) if f.endswith('.xml') and not f.startswith('FR') and len(f) == 13]

    # Get the current date and time for the filename
    current_time = datetime.now().strftime("%Y%m%d%H%M%S")

    # Define the output DOCX file path with the timestamp and desired name format
    docx_file = os.path.join(folder_path, f'output_znieff_xml2docx_com_{current_time}.docx')

    # Create a new DOCX document
    doc = Document()

    # Process each XML file
    for xml_file in xml_files:
        # Construct the full path to the XML file
        full_path = os.path.join(folder_path, xml_file)

        # Call the function to process the XML and append to the DOCX
        xml_to_docx(full_path, doc)

    # Clean up the document by removing double spaces and line breaks
    clean_document(doc)

    # Save the document after processing all XML files
    doc.save(docx_file)
    print(f"Document successfully created: {docx_file}")

# Demander à l'utilisateur de spécifier le chemin du dossier contenant les fichiers XML
folder_path = obtain_folder_path()

# Appeler la fonction avec le chemin du dossier fourni par l'utilisateur
process_xml_files_in_folder(folder_path)
