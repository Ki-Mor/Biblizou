"""
Author : ExEco Environnement
Edition date : 2025/02
Name : 22_n2000_xml2docx_desc
Group : Biblio_PatNat
"""

import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import os
import sys

def obtain_folder_path():
    return sys.argv[1] if len(sys.argv) > 1 else input("Entrez le chemin du dossier contenant les fichiers XML : ")

def xml_to_docx(xml_file, doc):
    """
    Parse the XML file and append the concatenated <SITE_NAME>, <SITECODE>,
    followed by a line break and the content inside <QUALITY> and <VULNAR> to the DOCX.
    """
    try:
        # Parse the XML file
        tree = ET.parse(xml_file)
        root = tree.getroot()

        # Iterate through <BIOTOP> elements
        for n2000_elem in root.iter('BIOTOP'):
            # Extract SITECODE and SITE_NAME from the current <BIOTOP> element
            sitecode_elem = n2000_elem.find('SITECODE')
            site_name_elem = n2000_elem.find('SITE_NAME')

            # Initialize variables for SITE_NAME and SITECODE
            site_name = site_name_elem.text if site_name_elem is not None else ''
            sitecode = sitecode_elem.text if sitecode_elem is not None else ''

            # Concatenate SITE_NAME and SITECODE
            combined_text = f"{site_name} - {sitecode}"

            # Add the concatenated text as a paragraph in the DOCX document (with formatting)
            para = doc.add_paragraph()
            run = para.add_run(combined_text)
            run.bold = True
            run.underline = True
            run.font.color.rgb = RGBColor(0, 153, 153)  # Color #009999
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

            # Now, we need to find <QUALITY> and <VULNAR> inside <COMMENTAIRE_ROW> elements within <COMMENTAIRE>
            commentaire_elem = n2000_elem.find('COMMENTAIRE')
            if commentaire_elem is not None:
                for commentaire_row_elem in commentaire_elem.iter('COMMENTAIRE_ROW'):
                    # Extract <QUALITY> and <VULNAR>
                    quality_elem = commentaire_row_elem.find('QUALITY')
                    vulnar_elem = commentaire_row_elem.find('VULNAR')

                    # Add the content of <QUALITY> if it exists
                    if quality_elem is not None and quality_elem.text:
                        quality_text = quality_elem.text.strip()
                        if quality_text:  # Make sure there is actual text to add
                            para = doc.add_paragraph(quality_text)
                            para.paragraph_format.left_indent = Pt(28)  # 1 cm tabulation
                            run = para.runs[0]
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

                    # Add the content of <VULNAR> if it exists
                    if vulnar_elem is not None and vulnar_elem.text:
                        vulnar_text = vulnar_elem.text.strip()
                        if vulnar_text:  # Make sure there is actual text to add
                            para = doc.add_paragraph(vulnar_text)
                            para.paragraph_format.left_indent = Pt(28)  # 1 cm tabulation
                            run = para.runs[0]
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

            # Add a line break between entries (if needed)
            doc.add_paragraph('')

    except ET.ParseError as e:
        print(f"Error parsing XML file {xml_file}: {e}")
    except Exception as e:
        print(f"Unexpected error with file {xml_file}: {e}")

def clean_document(doc):
    """
    Remove double spaces and double line breaks from the document, preserving the formatting.
    """
    for para in doc.paragraphs:
        # Iterate over the runs in the paragraph
        for run in para.runs:
            # Clean the text of the current run (remove double spaces)
            run.text = ' '.join(run.text.split())

        # Remove unnecessary empty paragraphs (double line breaks)
        if para.text.strip() == '':
            para.clear()  # Clears the paragraph content, removing extra empty lines

def process_xml_files_in_folder(folder_path):
    """
    Process all XML files in a folder and append their extracted data into a single DOCX file.
    """
    # Validate the folder path
    if not os.path.isdir(folder_path):
        print(f"The path {folder_path} is not a valid directory.")
        return

    # List all XML files in the folder
    xml_files = [f for f in os.listdir(folder_path) if f.startswith('FR') and f.endswith('.xml') and len(f) == 13]

    # If no XML files are found, exit early
    if not xml_files:
        print("No XML files found in the folder.")
        return

    # Get the current date and time for the filename
    current_time = datetime.now().strftime("%Y%m%d%H%M%S")

    # Define the output DOCX file path with the timestamp and desired name format
    docx_file = os.path.join(folder_path, f'N2000_Descriptions_des_sites_{current_time}.docx')

    # Create a new DOCX document
    doc = Document()

    # Process each XML file
    for xml_file in xml_files:
        # Construct the full path to the XML file
        full_path = os.path.join(folder_path, xml_file)

        # Call the function to process the XML and append to the DOCX
        xml_to_docx(full_path, doc)

    # Clean up the document by removing double spaces and line breaks
    clean_document(doc)

    try:
        # Save the document after processing all XML files
        doc.save(docx_file)
        print(f"Document successfully created: {docx_file}")
    except Exception as e:
        print(f"Error saving document: {e}")

# Demander à l'utilisateur de spécifier le chemin du dossier contenant les fichiers XML
folder_path = obtain_folder_path()

# Appeler la fonction avec le chemin du dossier fourni par l'utilisateur
process_xml_files_in_folder(folder_path)
