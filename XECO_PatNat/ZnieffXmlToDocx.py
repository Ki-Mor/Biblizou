"""
Auteur : ExEco Environnement - François Botcazou
Date de création : 2025/02
Dernière mise à jour : 2025/03
Version : 1.0
Nom : ZnieffXmlToDocx.py
Groupe : Biblizou_PatNat
Description : Module pour extraire les descriptions des sites ZNIEFF des fichiers XML ZNIEFF et créer un fichier de synthèse docx.
    Il est conçu pour être utilisé dans une extension QGIS
Dépendances :
    - Python 3.x
    - QGIS (QgsMessageLog)
    - xml.etree.ElementTree
    - python-docx
    - os, datetime

Utilisation :
    Ce module doit être appelé depuis une extension QGIS. Il prend en entrée un
    dossier contenant des fichiers XML et génère un fichier DOCX récapitulatif."""

import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import os
from qgis.core import QgsMessageLog, Qgis


class ZnieffXmlToDocx:
    def __init__(self, folder_path):
        self.folder_path = folder_path
        self.doc = Document()

    def log(self, message, level=Qgis.Info):
        QgsMessageLog.logMessage(message, 'ZNIEFF_XML2DOCX', level)

    def xml_to_docx(self, xml_file):
        try:
            tree = ET.parse(xml_file)
            root = tree.getroot()
            for znieff_elem in root.iter('ZNIEFF'):
                lb_zn = znieff_elem.findtext('LB_ZN', '')
                nm_sffzn = znieff_elem.findtext('NM_SFFZN', '')
                tx_gene_elem = znieff_elem.find('TX_GENE')
                combined_text = f"{lb_zn} - {nm_sffzn}"
                para = self.doc.add_paragraph()
                run = para.add_run(combined_text)
                run.bold = True
                run.underline = True
                run.font.color.rgb = RGBColor(0, 153, 153)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                if tx_gene_elem is not None:
                    for p_elem in tx_gene_elem.findall('p'):
                        p_text = p_elem.text or ''
                        para = self.doc.add_paragraph(p_text)
                        para.paragraph_format.left_indent = Pt(28)
                        run = para.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                self.doc.add_paragraph('')
        except ET.ParseError as e:
            self.log(f"Error parsing XML file {xml_file}: {e}", Qgis.Warning)
        except Exception as e:
            self.log(f"Unexpected error with file {xml_file}: {e}", Qgis.Critical)

    def clean_document(self):
        for para in self.doc.paragraphs:
            for run in para.runs:
                run.text = ' '.join(run.text.split())
            if para.text.strip() == '':
                para.clear()

    def process_folder(self):
        if not os.path.isdir(self.folder_path):
            self.log(f"Invalid directory: {self.folder_path}", Qgis.Warning)
            return None
        xml_files = [f for f in os.listdir(self.folder_path) if
                     f.endswith('.xml') and not f.startswith('FR') and len(f) == 13]
        if not xml_files:
            self.log("No valid XML files found.", Qgis.Info)
            return None
        current_time = datetime.now().strftime("%Y%m%d%H%M%S")
        docx_file = os.path.join(self.folder_path, f'ZNIEFF_Descriptions_des_sites_{current_time}.docx')
        for xml_file in xml_files:
            self.xml_to_docx(os.path.join(self.folder_path, xml_file))
        self.clean_document()
        self.doc.save(docx_file)
        self.log(f"Document successfully created: {docx_file}", Qgis.Success)
        return docx_file
